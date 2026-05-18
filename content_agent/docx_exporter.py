"""
Markdown → 精美 Word 文档排版

支持标题、列表、代码块、引用、行内格式等 Markdown 元素的高质量转换。
"""

import os
import re
import tempfile
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _set_run_font(run, font_name="Arial", font_size=None, bold=False, color=None):
    """
    设置 run 字体，同时设置东亚字体回退避免中文显示异常
    """
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    if font_size:
        run.font.size = font_size
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def _add_shading(run, fill_color: str):
    """给 run 添加背景色"""
    rPr = run._element.get_or_add_rPr()
    shd = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
            fill_color,
        )
    )
    rPr.append(shd)


def _parse_inline(paragraph, text: str):
    """解析行内 Markdown 格式：加粗 **、斜体 *、行内代码 `、链接 []()"""
    if not text:
        return

    # 正则分割：匹配加粗、斜体、行内代码、链接
    pattern = r'(\*\*[^*]+?\*\*|\*[^*]+?\*|_[^_]+?_|`[^`]+?`|\[[^\]]+?\]\([^)]+?\))'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            _set_run_font(run)
            run.italic = True
        elif part.startswith("_") and part.endswith("_") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            _set_run_font(run)
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            _set_run_font(run, font_name="Courier New", font_size=Pt(9.5))
            _add_shading(run, "F0F0F0")
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            m = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if m:
                run = paragraph.add_run(m.group(1))
                _set_run_font(run, color=RGBColor(0x00, 0x66, 0xCC))
                run.underline = True
            else:
                paragraph.add_run(part)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run)


def _add_code_block(doc, lines: list[str]):
    """添加代码块：用表格容器+浅灰背景+等宽字体"""
    # 移除开头的 ```language 如果有
    code_content = "\n".join(lines).strip()
    if code_content.startswith("```"):
        code_content = code_content[code_content.find("\n") + 1 :]
    if code_content.endswith("```"):
        code_content = code_content[: code_content.rfind("```")]
    code_content = code_content.strip("\n")

    if not code_content:
        return

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)

    # 单元格背景色
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(
        parse_xml(
            r'<w:shd {} w:fill="F8F8F8" w:val="clear"/>'.format(
                'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            )
        )
    )

    # 清空默认段落，重新添加
    p = cell.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15

    run = p.add_run(code_content)
    _set_run_font(
        run,
        font_name="Courier New",
        font_size=Pt(9),
        color=RGBColor(0x33, 0x33, 0x33),
    )

    # 单元格内边距
    tbl_pr = table._tblPr
    tbl_pr.append(
        parse_xml(
            r'<w:tblCellMar {}>'
            r'<w:top w:w="80" w:type="dxa"/>'
            r'<w:left w:w="120" w:type="dxa"/>'
            r'<w:bottom w:w="80" w:type="dxa"/>'
            r'<w:right w:w="120" w:type="dxa"/>'
            r"</w:tblCellMar>".format(
                'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            )
        )
    )

    # 段后留白
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def render_markdown_to_docx(content: str, title: str = None) -> str:
    """
    将 Markdown 内容渲染为精美排版的 Word 文档。

    Args:
        content: Markdown 格式的文本
        title: 文档标题（可选）

    Returns:
        临时 .docx 文件路径

    Raises:
        ImportError: 未安装 python-docx
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx 未安装，请运行: pip install python-docx"
        )

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ── 默认段落样式 ──
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # ── 文档标题 ──
    if title:
        p = doc.add_heading(level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        _set_run_font(run, font_size=Pt(20), bold=True)
        p.paragraph_format.space_after = Pt(18)
        doc.add_paragraph()

    # ── 逐行解析 Markdown ──
    lines = content.split("\n")
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块边界
        if stripped.startswith("```"):
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # 标题
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            p = doc.add_heading(level=min(level, 3))
            run = p.add_run(text)
            size_map = {1: Pt(18), 2: Pt(15), 3: Pt(13)}
            _set_run_font(run, font_size=size_map.get(level, Pt(12)), bold=True)
            p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # 无序列表
        ul_match = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if ul_match:
            text = ul_match.group(2)
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline(p, text)
            p.paragraph_format.left_indent = Cm(0.74)
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ol_match:
            text = ol_match.group(2)
            p = doc.add_paragraph(style="List Number")
            _parse_inline(p, text)
            p.paragraph_format.left_indent = Cm(0.74)
            i += 1
            continue

        # 引用块
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            _set_run_font(run, color=RGBColor(0x66, 0x66, 0x66))
            run.italic = True
            i += 1
            continue

        # 分隔线
        if re.match(r'^[-=_*]{3,}\s*$', stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run("\u2500" * 28)
            _set_run_font(run, color=RGBColor(0xCC, 0xCC, 0xCC), font_size=Pt(9))
            i += 1
            continue

        # 普通段落
        if stripped:
            p = doc.add_paragraph()
            _parse_inline(p, stripped)
            p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进约 2 字
        else:
            # 空行：控制段间距
            doc.add_paragraph()

        i += 1

    # ── 页脚 ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(
        f"由 Content Agent 生成 · {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    _set_run_font(run, font_size=Pt(8), color=RGBColor(0x99, 0x99, 0x99))

    # 保存为临时文件
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path
